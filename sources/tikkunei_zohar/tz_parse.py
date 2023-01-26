from docx import Document
from bs4 import BeautifulSoup
from docx2python import docx2python
from base64 import b64encode
from tz_base import *
import re
import csv

import logging

class TzParser(object):
    PUNCTUATION = {
        # '‘': True,
        '.': True,
        ':': True
    }
    def __init__(self, filename, volume, language="english", starting_tikkun=None, starting_daf=None):
        self.file = filename
        self.volume = volume

        self.doc_rep = None
        self.doc_rep_he = None
        self.elem_cursor = None
        self.processed_elem_cursor = None
        self.word_cursor = None
        self.language = language

        self.tikkun_number = count(0)
        self.paragraph_number = count(0)

        self.words = []
        self.phrases = []
        self.lines = []
        self.paragraphs = []
        self.dapim = []
        self.tikkunim = []
        self.quoted = []

        self.he_words = []
        self.he_lines = []
        self.he_paragraphs = []
        self.he_dapim = []
        self.he_tikkunim = []

        self.word = None
        self.phrase = None
        self.line = None
        self.paragraph = None
        self.daf = starting_daf if starting_daf else None
        self.tikkun = starting_tikkun if starting_tikkun else None
        self.append_to_previous = False
        self.current_footnote = None
        self.parsing_footnote = False
        self.previous_citation = None

        self.title = None

    def read_file(self):
        self.doc_rep = self.get_document_representation()
        self.doc_rep_he = self.get_document_representation_he()
        self.parse_contents()

    def get_document_representation(self):
        """Returns abstracted representation of doc"""
        pass

    def get_document_representation_he(self):
        pass

    def parse_contents(self):
        """Parse the document"""
        self.title = self.get_title()
        # if self.language == "hebrew":
        #     self.parse_hebrew_contents()
        # else:  # parse english
        self.elem_cursor = self.move_cursor()
        while self.elem_cursor:
            self.process_cursor()
            self.elem_cursor = self.move_cursor()
        if self.language == "bi":
            self.parse_hebrew_contents()

    def parse_hebrew_contents(self):
        pass

    def process_cursor(self):
        self.processed_elem_cursor = self.get_processed_elem(self.elem_cursor)
        #print(self.processed_elem_cursor['class'])
        if self.cursor_is_daf():
            self.daf = self.get_daf()
        elif self.cursor_is_tikkun():
            self.tikkun = self.get_tikkun()
        elif self.cursor_is_paragraph():
            if self.tikkun:
                self.process_paragraph()
        elif self.cursor_contains_words():  # this is not used for HtmlParser
            self.process_words()
            # self.word = self.get_next_word()
            # while self.word:
            #     self.word = self.get_next_word()
            # self.word_cursor = None

    def create_new_paragraph(self):
        self.paragraph = Paragraph(self.tikkun, self.daf, next(self.daf.paragraph_number))
        self.paragraphs.append(self.paragraph)
        self.tikkun.paragraphs.append(self.paragraph)
        self.daf.paragraphs.append(self.paragraph)
        self.line = self.paragraph.add_new_line()
        self.lines.append(self.line)
        self.append_to_previous = False

    def get_processed_elem(self):
        pass

    def cursor_is_daf(self):
        pass

    def get_daf(self):
        pass

    def cursor_is_tikkun(self):
        pass

    def get_tikkun(self):
        pass

    def cursor_is_paragraph(self):
        pass

    def process_paragraph(self):
        pass

    def cursor_contains_words(self):
        pass

    def get_next_word(self):
        """Processes the next word in the cursor element"""
        pass

    def move_cursor(self):
        pass

    def get_title(self):
        """get the title of the document"""
        pass


class DocsTzParser(TzParser):
    def __init__(self, filename, volume, language="english", starting_tikkun=None, starting_daf=None):
        TzParser.__init__(self, filename, volume, language, starting_tikkun, starting_daf)
        self.paragraph_index = 0

    def get_document_representation(self):
        return Document(self.file)

    def get_document_representation_he(self):
        filename = self.file.replace('.docx', '_he.docx')
        return Document(filename)

    def get_title(self):
        return self.file

    def move_cursor(self):
        if self.paragraph_index < len(self.doc_rep.paragraphs):
            cursor = self.doc_rep.paragraphs[self.paragraph_index]
            self.paragraph_index += 1
            # print(cursor.text)
            return cursor
        else:
            return None

    @staticmethod
    def get_processed_elem(elem_cursor):
        return elem_cursor

    def cursor_is_daf(self):
        return re.search(r'\[[0-9]+[ab]\]', self.processed_elem_cursor.text) is not None

    def get_daf(self):
        stripped_text = self.processed_elem_cursor.text.strip()
        stripped_text = stripped_text.lstrip("[")
        stripped_text = stripped_text.rstrip("]")
        daf = Daf(stripped_text)
        self.dapim.append(daf)
        return daf

    def cursor_is_tikkun(self):
        return re.match(r'^\[?Tiqun', self.processed_elem_cursor.text) is not None or \
               re.match(r'^\[?The Introduction of Tiqunei', self.processed_elem_cursor.text) is not None

    def get_tikkun(self):
        # print(self.processed_elem_cursor.text)
        tikkun = Tikkun(self.processed_elem_cursor.text, next(self.tikkun_number))
        for run in self.processed_elem_cursor.runs:
            if run.text == '' and run.element.style == 'EndnoteReference':
                tikkun.words.append(Word("[ENDNOTE_REFERENCE]", None, None, self.daf, tikkun, None))
            else:
                tikkun.words.append(Word(run.text, None, None, self.daf, tikkun, None))
        print([word.text for word in tikkun.words])
        self.tikkunim.append(tikkun)
        return tikkun # TODO: handle tikkun footnotes

    def cursor_is_paragraph(self):
        return self.processed_elem_cursor.text.strip() == '' and self.processed_elem_cursor.style.name == 'Normal'

    def process_paragraph(self):
        if not self.paragraph or len(self.paragraph.words) > 0:
            self.create_new_paragraph()
    # def cursor_contains_words(self):
    #     if paragraph.text

    def cursor_contains_words(self):
        return True  # catchall????

    @staticmethod
    def get_formatting(run):
        if run.bold and run.italic:
            formatting = Formatting.BOLD_ITALICS
        elif run.bold:
            formatting = Formatting.BOLD
        elif run.italic:
            formatting = Formatting.ITALICS
        elif run.font.color.rgb is not None:
            formatting = Formatting.FADED
        else:
            formatting = None
        return formatting

    def parse_hebrew_contents(self):
        daf_index = 0
        tikkun_index = 0
        paragraph_index = 0
        for paragraph in self.doc_rep_he.paragraphs: #daf
            if re.search(r'[\u0590-\u05fe]{1,3}/[אב]', paragraph.text):
                daf = re.search(r'[\u0590-\u05fe]{1,3}/[אב]', paragraph.text).group(0)
                self.daf = self.dapim[daf_index]
                self.daf.he_name = daf
                print(self.daf.name)
                print(self.daf.he_name)
                daf_index += 1
                paragraph_index = 0
            elif re.match(r'תּקונא', paragraph.text) or re.match(r'תקונא', paragraph.text):
                self.tikkun = self.tikkunim[tikkun_index]
                self.tikkun.he_name = paragraph.text
                tikkun_index +=1
            elif paragraph.text != '':
                if paragraph_index < len(self.daf.paragraphs):
                    self.paragraph = self.daf.paragraphs[paragraph_index]
                else:  # somehow misaligned??
                    self.paragraph = Paragraph(self.tikkun, self.daf, next(self.paragraph_number))
                    self.daf.paragraphs.append(self.paragraph)
                self.paragraph.he_words = self.clean_he(paragraph.text)
                paragraph_index += 1

    @staticmethod
    def clean_he(hebrew_text):
        return hebrew_text.replace("\n", "</br>")

    def process_words(self):
        if self.line is None:
            self.create_new_paragraph()
        for run in self.processed_elem_cursor.runs:
            if self.parsing_footnote:
                for char in run.text:
                    self.process_word(char, True, run)
            else:
                if len(run.text) > 0 and run.text[0] == ' ':
                    self.append_to_previous = False
                # if run.style == Normal??
                # for i, elem_word in enumerate(run.text.split(' ')):
                # else process footnote?
                # append to previous?????
                if not self.append_to_previous or len(self.line.phrases) == 0:  # new phrase for run
                    formatting = DocsTzParser.get_formatting(run)
                    self.phrase = self.line.add_new_phrase(formatting)
                    self.phrases.append(self.phrase)
                if self.current_footnote and self.phrase.formatting and not self.parsing_footnote:
                    self.current_footnote.anchor = self.phrase
                    self.phrase.footnotes.append(self.current_footnote)
                    self.current_footnote = None
                if run.text == ' ':
                    self.append_to_previous = False
                else:
                    for i, elem_word in enumerate(run.text.split(' ')):
                        if i > 0:
                            self.append_to_previous = False
                        if '{' in elem_word or '}' in elem_word or '\n' in elem_word or '‹' in elem_word or '›' in elem_word:
                            for char in elem_word:
                                self.process_word(char, True)
                        elif elem_word == '':
                            # parse w/ beautiful soup
                            pass
                        else:
                            self.process_word(elem_word)
                    # print(run.text)
                            if len(run.text) > 0 and run.text[-1] != ' ':
                                self.append_to_previous = True
                            else:
                                self.append_to_previous = False

    def process_word(self, word, char=False, run=None):
        if word == '{':
            self.parsing_footnote = True
            self.append_to_previous = False
            self.current_footnote = Footnote(FootnoteType.CITATION, None)
        elif word == '}':
            self.parsing_footnote = False
            self.append_to_previous = False
        elif word == '‹':
            self.parsing_footnote = True
            self.append_to_previous = False
            self.current_footnote = Footnote(FootnoteType.SYMBOL, None)  # TODO: Get specific type from BS
        elif word == '›':
            self.parsing_footnote = False
            if self.current_footnote:  # TODO: handle errors
                self.current_footnote.anchor = self.word
                self.word.footnotes.append(self.current_footnote)
                self.current_footnote = None
        elif word == '\n':
            self.line = self.paragraph.add_new_line()
            self.lines.append(self.line)
            self.append_to_previous = False

        else:
            if self.parsing_footnote:
                if char: #why
                    self.current_footnote.text += word
                    self.current_footnote.formatting = DocsTzParser.get_formatting(run) if run else \
                        self.current_footnote.formatting
            elif self.append_to_previous:
                if self.parsing_footnote:
                    self.current_footnote.text += word
                else:
                    self.word.add_to_word(word)
                # self.append_to_previous = False
            else:
                if self.parsing_footnote:  # delete? this won't ever get hit
                    self.current_footnote.text += " "
                    self.current_footnote.text += word
                else:
                    self.word = self.phrase.add_new_word(word)
                    self.word.footnotes = []
                    self.words.append(self.word)
                if self.current_footnote and not self.parsing_footnote:
                    self.current_footnote.anchor = self.word
                    self.word.footnotes.append(self.current_footnote)
                    self.current_footnote = None
            #print(self.word.text)


class DocsTzParser2(TzParser):
    def __init__(self, filename, volume, language="english", starting_tikkun=None, starting_daf=None):
        TzParser.__init__(self, filename, volume, language, starting_tikkun, starting_daf)
        self.paragraph_index = 0

    def get_document_representation(self):
        return docx2python(self.file, html=True)  # YOU ARE HERE

    def get_document_representation_he(self):
        filename = self.file.replace('.docx', '_he.docx')
        return Document(filename)

    def get_title(self):
        return self.file

    def move_cursor(self):
        if self.paragraph_index < len(self.doc_rep.document[0][0][0]):
            cursor = self.doc_rep.document[0][0][0][self.paragraph_index]
            self.paragraph_index += 1
            # print(cursor.text)
            return cursor
        else:
            return None

    @staticmethod
    def get_processed_elem(elem_cursor):
        return BeautifulSoup(elem_cursor, 'html.parser')

    def cursor_is_daf(self):
        return re.search(r'\[[0-9]+[ab]\]', self.processed_elem_cursor.text) is not None

    def get_daf(self):
        stripped_text = self.processed_elem_cursor.text.strip()
        stripped_text = stripped_text.lstrip("[")
        stripped_text = stripped_text.rstrip("]")
        daf = Daf(stripped_text)
        self.dapim.append(daf)
        return daf

    def cursor_is_tikkun(self):
        return re.match(r'^\[?Tiqun', self.processed_elem_cursor.text) is not None or \
               re.match(r'^\[?The Introduction of Tiqunei', self.processed_elem_cursor.text) is not None

    def process_tikkun_name_contents(self, tikkun, segment):
        if isinstance(segment, str):
            if segment.startswith("----endnote"):
                endnote_number = segment.lstrip("----endnote")[0]
                tikkun.words.append(Word("[ENDNOTE_REFERENCE_" + endnote_number + "]", None, None, self.daf, tikkun, None))
            else:
                tikkun.words.append(Word(segment, None, None, self.daf, tikkun, None))
        else:
            for seg_child in segment.children:
                self.process_tikkun_name_contents(seg_child)

    def get_tikkun(self):
        # print(self.processed_elem_cursor.text)
        tikkun = Tikkun(self.processed_elem_cursor.text, next(self.tikkun_number))
        for segment in self.processed_elem_cursor:
            self.process_tikkun_name_contents(tikkun, segment)
        #     if run.text == '' and run.element.style == 'EndnoteReference':
        #         tikkun.words.append(Word("[ENDNOTE_REFERENCE]", None, None, self.daf, tikkun, None))
        #     else:
        #         tikkun.words.append(Word(run.text, None, None, self.daf, tikkun, None))
        # print([word.text for word in tikkun.words])
        # self.tikkunim.append(tikkun)
        return tikkun # TODO: handle tikkun footnotes

    def cursor_is_paragraph(self):
        return self.processed_elem_cursor.text.strip() == '' and self.processed_elem_cursor.style.name == 'Normal'

    def process_paragraph(self):
        if not self.paragraph or len(self.paragraph.words) > 0:
            self.create_new_paragraph()
    # def cursor_contains_words(self):
    #     if paragraph.text

    def cursor_contains_words(self):
        return True  # catchall????

    @staticmethod
    def get_formatting(run):
        if run.bold and run.italic:
            formatting = Formatting.BOLD_ITALICS
        elif run.bold:
            formatting = Formatting.BOLD
        elif run.italic:
            formatting = Formatting.ITALICS
        elif run.font.color.rgb is not None:
            formatting = Formatting.FADED
        else:
            formatting = None
        return formatting

    def parse_hebrew_contents(self):
        daf_index = 0
        tikkun_index = 0
        paragraph_index = 0
        for paragraph in self.doc_rep_he.paragraphs: #daf
            if re.search(r'[\u0590-\u05fe]{1,3}/[אב]', paragraph.text):
                daf = re.search(r'[\u0590-\u05fe]{1,3}/[אב]', paragraph.text).group(0)
                self.daf = self.dapim[daf_index]
                self.daf.he_name = daf
                print(self.daf.name)
                print(self.daf.he_name)
                daf_index += 1
                paragraph_index = 0
            elif re.match(r'תּקונא', paragraph.text) or re.match(r'תקונא', paragraph.text):
                self.tikkun = self.tikkunim[tikkun_index]
                self.tikkun.he_name = paragraph.text
                tikkun_index +=1
            elif paragraph.text != '':
                if paragraph_index < len(self.daf.paragraphs):
                    self.paragraph = self.daf.paragraphs[paragraph_index]
                else:  # somehow misaligned??
                    self.paragraph = Paragraph(self.tikkun, self.daf, next(self.paragraph_number))
                    self.daf.paragraphs.append(self.paragraph)
                self.paragraph.he_words = self.clean_he(paragraph.text)
                paragraph_index += 1

    @staticmethod
    def clean_he(hebrew_text):
        return hebrew_text.replace("\n", "</br>")

    def process_words(self):
        if self.line is None:
            self.create_new_paragraph()
        for run in self.processed_elem_cursor.runs:
            if self.parsing_footnote:
                for char in run.text:
                    self.process_word(char, True, run)
            else:
                if len(run.text) > 0 and run.text[0] == ' ':
                    self.append_to_previous = False
                # if run.style == Normal??
                # for i, elem_word in enumerate(run.text.split(' ')):
                # else process footnote?
                # append to previous?????
                if not self.append_to_previous or len(self.line.phrases) == 0:  # new phrase for run
                    formatting = DocsTzParser.get_formatting(run)
                    self.phrase = self.line.add_new_phrase(formatting)
                    self.phrases.append(self.phrase)
                if self.current_footnote and self.phrase.formatting and not self.parsing_footnote:
                    self.current_footnote.anchor = self.phrase
                    self.phrase.footnotes.append(self.current_footnote)
                    self.current_footnote = None
                if run.text == ' ':
                    self.append_to_previous = False
                else:
                    for i, elem_word in enumerate(run.text.split(' ')):
                        if i > 0:
                            self.append_to_previous = False
                        if '{' in elem_word or '}' in elem_word or '\n' in elem_word or '‹' in elem_word or '›' in elem_word:
                            for char in elem_word:
                                self.process_word(char, True)
                        elif elem_word == '':
                            # parse w/ beautiful soup
                            pass
                        else:
                            self.process_word(elem_word)
                    # print(run.text)
                            if len(run.text) > 0 and run.text[-1] != ' ':
                                self.append_to_previous = True
                            else:
                                self.append_to_previous = False

    def process_word(self, word, char=False, run=None):
        if word == '{':
            self.parsing_footnote = True
            self.append_to_previous = False
            self.current_footnote = Footnote(FootnoteType.CITATION, None)
        elif word == '}':
            self.parsing_footnote = False
            self.append_to_previous = False
        elif word == '‹':
            self.parsing_footnote = True
            self.append_to_previous = False
            self.current_footnote = Footnote(FootnoteType.SYMBOL, None)  # TODO: Get specific type from BS
        elif word == '›':
            self.parsing_footnote = False
            if self.current_footnote:  # TODO: handle errors
                self.current_footnote.anchor = self.word
                self.word.footnotes.append(self.current_footnote)
                self.current_footnote = None
        elif word == '\n':
            self.line = self.paragraph.add_new_line()
            self.lines.append(self.line)
            self.append_to_previous = False

        else:
            if self.parsing_footnote:
                if char: #why
                    self.current_footnote.text += word
                    self.current_footnote.formatting = DocsTzParser.get_formatting(run) if run else \
                        self.current_footnote.formatting
            elif self.append_to_previous:
                if self.parsing_footnote:
                    self.current_footnote.text += word
                else:
                    self.word.add_to_word(word)
                # self.append_to_previous = False
            else:
                if self.parsing_footnote:  # delete? this won't ever get hit
                    self.current_footnote.text += " "
                    self.current_footnote.text += word
                else:
                    self.word = self.phrase.add_new_word(word)
                    self.word.footnotes = []
                    self.words.append(self.word)
                if self.current_footnote and not self.parsing_footnote:
                    self.current_footnote.anchor = self.word
                    self.word.footnotes.append(self.current_footnote)
                    self.current_footnote = None
            #print(self.word.text)


class HtmlTzParser(TzParser):
    FOOTNOTES = {
        "CharOverride-1": FootnoteType.SYMBOL,
        "CharOverride-3": FootnoteType.SYMBOL,
        "CharOverride-4": FootnoteType.SYMBOL,
        "CharOverride-5": FootnoteType.SYMBOL,
        "CharOverride-6": FootnoteType.SYMBOL,
        "CharOverride-7": FootnoteType.SYMBOL,
        "CharOverride-8": FootnoteType.SYMBOL,
        "CharOverride-10": FootnoteType.SYMBOL,
        "FNR---verse-English": FootnoteType.FOOTNOTE,
        "er": FootnoteType.ENDNOTE,
        "stars":
            {"\uF0A4": FootnoteType.STAR,
             "": FootnoteType.DIAMONDS
             },
        "infinity": FootnoteType.INFINITY,
        "triangle": FootnoteType.TRIANGLE,
        "CharOverride-2": FootnoteType.SYMBOL,
    }

    FORMATTING_CLASSES = {
        "it-text": Formatting.ITALICS,
        "grey-text": Formatting.FADED,
        "bd-it-text": Formatting.BOLD_ITALICS,
        "bd": Formatting.BOLD,
    }

    SIDENOTE_CLASSES = {
        "it-text---side-notes": Formatting.ITALICS,
        "bd-it-text---side-notes": Formatting.BOLD_ITALICS,
    }

    APPEND_TO_WORD = [
        "dot-italic",
        'dot-roman',
    ]

    def __init__(self, filename, volume, language="english", starting_tikkun=None, starting_daf=None, endnotes_doc=None):
        TzParser.__init__(self, filename, volume, language, starting_tikkun, starting_daf)
        self.continue_phrase = False
        self.endnote_dict = self.get_endnote_dict(endnotes_doc)

    def get_endnote_dict(self, endnotes_doc):
        endnote_dict = {}
        endnote_dict['579'] = 'There is some confusion about the designation of Tiqun 15 as being the Ellul reading for ‘Day 8’, since it is also the designation found for the commencement of Tiqun 13 on TZ 27a.'
        if endnotes_doc: # assumes endnotes are doc file
            endnotes_doc_rep = Document(endnotes_doc)
            prev_endnote = '1'
            for paragraph in endnotes_doc_rep.paragraphs:
                if not paragraph.text.strip() == '':
                    if paragraph.text[0] == '0':
                        endnote_dict['2'] = paragraph.text.lstrip('0')
                    elif paragraph.text[0] == '1' and not paragraph.text[1].isnumeric():
                        endnote_dict['1'] = paragraph.text.lstrip('1')
                    elif paragraph.text.startswith('96') and '97' not in endnote_dict:
                        endnote_dict['97'] = paragraph.text.lstrip('96')
                    else:
                        endnote_number = re.match(r'[0-9]+', paragraph.text)
                        if endnote_number:
                            endnote_number = endnote_number[0]
                            if int(endnote_number) < 578:
                                endnote_number_real = int(endnote_number) + 1
                            else:
                                endnote_number_real = int(endnote_number) + 2
                            endnote_dict[str(endnote_number_real)] = paragraph.text.lstrip(endnote_number)
                            prev_endnote = str(endnote_number_real)
                        else:
                            endnote_dict[prev_endnote] += '</br>'
                            endnote_dict[prev_endnote] += paragraph.text
        return endnote_dict

    def get_document_representation(self):
        with open(self.file, 'r') as file:
            contents = file.read()
            return BeautifulSoup(contents, 'html.parser')

    def get_title(self):
        return self.doc_rep.title.getText()

    def move_cursor(self):
        """Return the next element at the line level"""
        if not self.elem_cursor:
            # if self.language == "english":
            cursor = self.doc_rep.find_all("div", class_="_idGenObjectStyleOverride-1")[0].contents[0]
            if cursor.name:
                return cursor
        else:
            cursor = self.elem_cursor
        # augment
        cursor = cursor.next_sibling
        while cursor and not cursor.name:
            cursor = cursor.next_sibling
        return cursor

    @staticmethod
    def get_processed_elem(elem_cursor):
        """returns the elem if it is not just a container, otherwise, return useful element"""
        if "_idGenObjectLayout-1" in elem_cursor["class"]:
            return elem_cursor.div
        else:
            return elem_cursor

    def cursor_is_daf(self):
        return "daf" in self.processed_elem_cursor["class"]

    def get_daf(self):
        daf = Daf(self.processed_elem_cursor.p.text)
        self.dapim.append(daf)
        return daf

    def cursor_is_tikkun(self):
        return "chapter-number-title" in self.processed_elem_cursor["class"]

    def add_tikkun_word_to_list(self, elems, arr):
        for child in elems:
            if isinstance(child, str):
                arr.append(child)
            elif not any (x in child.get('class', ['CharOverride-1']) for x in ['CharOverride-1', 'er---chapter-number-title', 'CharOverride-16']):
                self.add_tikkun_word_to_list(child, arr)

    def get_tikkun(self):
        # TODO: further cleaning of tikkunim
        res = []
        self.add_tikkun_word_to_list(self.processed_elem_cursor.children, res)
        title = (''.join(res))
        tikkun = Tikkun(title, next(self.tikkun_number))
        for text in res:
            tikkun.words.append(Word(text, None, None, self.daf, tikkun, None))
        self.tikkunim.append(tikkun)
        return tikkun

    # def cursor_is_footnote(self):
    #     return self.processed_elem_cursor["class"] in ["books"]
    #
    # def get_footnotes(self):
    #     return self.processed_elem_cursor["class"] in ["books"]

    def cursor_is_paragraph(self):
        return self.processed_elem_cursor.name == 'p' and self.processed_elem_cursor['class']

    def process_words(self, elem, formatting, is_img_tag=False):
        """Process words or partial words and add to existing or new Word"""
        if (not self.append_to_previous or elem[0].isspace()) and not self.continue_phrase:  # new phrase
            self.append_to_previous = False
            self.phrase = self.line.add_new_phrase(formatting)
            self.phrases.append(self.phrase)
            if self.current_footnote:
                self.current_footnote.anchor = self.phrase
                self.phrase.footnotes.append(self.current_footnote)
                self.current_footnote = None
        self.continue_phrase = False
        if is_img_tag:
            self.word = self.phrase.add_new_word(elem)
            # self.word.footnotes = []
            self.paragraph.add_to_quoted_if_in_quotes(self.word)
            self.words.append(self.word)
        else:
            for i, elem_word in enumerate(elem.split()):
                if i == 0 and (self.append_to_previous or re.match(r'[?.\].,:!]+$', str(elem_word))) and len(self.line.words) > 0:  # ending punctuation when not new line
                    self.word.add_to_word(elem_word)
                elif re.match(r'[?.\[\].,:!]*‘[?.\[\].,:!]*$', str(elem_word)): # punctuation with backtaick
                    if re.match(r'[?.\[\].,:!]*‘[?.\[\].,:!]*$', str(elem_word)) and len(self.paragraph.quoted_cursor) > 0:   # self.paragraph.inside_quotes:
                        self.word.add_to_word(elem_word)
                    else:
                        self.word = self.phrase.add_new_word(elem_word)
                        # self.word.footnotes = []
                        self.paragraph.add_to_quoted_if_in_quotes(self.word)
                        self.words.append(self.word)
                    # elif self.line.inside_quotes:
                    # else:
                    #     self.word = self.phrase.add_new_word(elem_word)
                    #     self.words.append(self.word)
                    #     self.line.add_to_quoted_if_necessary(self.word)
                else:
                    self.word = self.phrase.add_new_word(elem_word)
                    # self.word.footnotes = []
                    self.words.append(self.word)
                    self.paragraph.add_to_quoted_if_in_quotes(self.word)
                #self.line.add_to_quoted_if_necessary(self.word)
                if '‘' in elem_word or '’' in elem_word:
                    if re.match(r'^[?./‘\[\].,:!]+[a-zA-Z0-9]+.*$', str(elem_word)):  # starts with ‘
                        self.paragraph.add_new_quoted()  # create a Quote()
                        self.paragraph.add_to_quoted_if_in_quotes(self.word)
                        if re.match(r'.*[?.\[\].,:!]*’[?.\[\].,:!]*', str(elem_word)):  # word ends in punctuation including ‘
                            self.paragraph.commit_quoted()
                        # number += 1
                        # if number != 1:
                        #    print("help!")
                        # self.line.add_new_quoted()
                    elif re.match(r'^[a-zA-Z0-9]+[?.\[\].,:!]*’[?.\[\].,:!]*$', str(elem_word)):  # word ends in punctuation including backtick but doesn't start with
                        self.paragraph.commit_quoted()
                    elif re.match(r'[A-Za-z]+’[A-Za-z]+', str(elem_word)):
                        pass
                    else:  # just ‘
                        # if self.line.inside_quotes:
                        #     self.line.commit_quoted()
                        # else:
                        #     self.line.add_new_quoted()
                        #     self.line.add_to_quoted(self.word)
                        # pass
                        # print(word)
                        if '‘' in elem_word:  # new quoted
                            if elem_word != '‘':
                                print(elem_word)
                            self.paragraph.add_new_quoted()
                            self.paragraph.add_to_quoted_if_in_quotes(self.word)
                        else:
                            if self.file == 'vol2.html' and self.has_vol_1_exceptions(elem_word):
                                pass
                            else:
                                if len(self.paragraph.quoted_cursor) == 0: #assume typo? # need to fix this see "and who raises her to her place"
                                    pass
                                    # print(self.word.text)
                                    # self.paragraph.add_new_quoted()
                                    # self.paragraph.add_to_quoted_if_in_quotes(self.word)
                                else:
                                    self.paragraph.commit_quoted()
                            # if len(self.paragraph.quoted_cursor) == 0:  # self.paragraph.inside_quotes:
                            #     print(elem_word)
                            #print(str(elem_word))

                    # if it's start
                    # if it's end
                    # if it's both

            self.append_to_previous = not elem[-1].isspace()

    def has_vol_1_exceptions(self, elem_word):
        if len(self.paragraphs) == 33 and len(self.paragraph.words) == 24:  # specific exception for typo
            return True
        elif '-’A' in elem_word:
            return True
        elif elem_word == '’ezer':
            return True
        elif elem_word == 'te-ru’ah':
            return True
        else:
            return False

    def process_footnote_material(self, elem):
        if not self.parsing_footnote or self.current_footnote.footnote_type == FootnoteType.SYMBOL:  # open a footnote or symbol is useless
            for footnote_type_class in HtmlTzParser.FOOTNOTES:
                if footnote_type_class in elem['class']:
                    footnote_type = HtmlTzParser.FOOTNOTES[footnote_type_class]
                    if not isinstance(footnote_type, Enum):
                        try:
                            footnote_type = HtmlTzParser.FOOTNOTES[footnote_type_class][elem.get_text()]
                        except KeyError as e:
                            footnote_type = FootnoteType.DIAMONDS
                            print('FootnoteType parsing anomaly detected: Daf ' + self.daf.name +
                                  '. Paragraph: ' + str(self.paragraph.paragraph_number))
                    break
            # else:
            #     footnote_type = FootnoteType.SYMBOL
            #    print("unexpected footnote!")
            # if "stars" in elem['class']:
            #     footnote_type = HtmlTzParser.FOOTNOTES["stars"]
            # elif "infinity" in elem['class']:
            #     footnote_type = HtmlTzParser.FOOTNOTES["infinity"]
            # elif "triangle" in elem['class']:
            #     footnote_type = HtmlTzParser.FOOTNOTES["triangle"]
            # elif "CharOverride-2" in elem['class']:
            #     footnote_type = HtmlTzParser.FOOTNOTES["CharOverride-2"]
            # elif "FNR---verse-English" in elem['class']:
            #     footnote_type = HtmlTzParser.FOOTNOTES['FNR---verse-English']
            # else:
            #     footnote_type = FootnoteType.SYMBOL
            format_class = HtmlTzParser.FORMATTING_CLASSES[[x for x in elem['class'] if x in HtmlTzParser.FORMATTING_CLASSES][0]] if \
                any(x in HtmlTzParser.FORMATTING_CLASSES for x in elem['class']) else None
            if not self.parsing_footnote:
                self.current_footnote = Footnote(footnote_type, format_class)
            else:
                self.current_footnote.footnote_type = footnote_type
        if self.current_footnote.footnote_type == HtmlTzParser.FOOTNOTES["FNR---verse-English"]:  # footnotes
            footnote_id = elem.contents[0].get('id').replace('-backlink', '')
            footnote_number = elem.text
            footnote_text = self.doc_rep.find("span", {"id": footnote_id}).text.lstrip(footnote_number).strip()
            self.current_footnote.anchor = self.word
            self.current_footnote.text = footnote_text
            self.current_footnote.footnote_number = footnote_number
            self.word.footnotes.append(self.current_footnote)
            self.current_footnote = None
            self.parsing_footnote = False
        elif self.current_footnote.footnote_type == HtmlTzParser.FOOTNOTES["er"]:  # endnote
            endnote_number = elem.text.strip()
            if endnote_number in self.endnote_dict:
                self.current_footnote.text = self.endnote_dict[endnote_number]
            else:
                print('missing endnote ' + endnote_number)
            self.current_footnote.anchor = self.word
            self.current_footnote.footnote_number = endnote_number
            self.word.footnotes.append(self.current_footnote)
            self.current_footnote = None
            self.parsing_footnote = False
        else:
            for letter in elem.text:
                self.process_footnote_text(letter)
            # if footnote_type is not HtmlTzParser.FOOTNOTES["CharOverride-2"]:  # not a book
            #     pass
            # else:
            #     self.current_footnote = Footnote(FootnoteType.CITATION, None)
            # if elem.text != "‹":
            #     print(elem.text)
        # TODO: Handle endnotes and footnotes in tikkunim titles
    def process_footnote_text(self, letter): # only run this if inside footnote
        if letter == '{':
            self.parsing_footnote = True
            self.append_to_previous = False
            self.current_footnote.footnote_type = FootnoteType.CITATION
        elif letter == '}': # end footnote
            if self.current_footnote.text == 'ibid.' or self.current_footnote.text == 'ibid':
                self.current_footnote.text = self.previous_citation
            elif self.current_footnote.footnote_type == FootnoteType.CITATION: # this will always be true so remove?
                self.previous_citation = self.current_footnote.text
            self.paragraph.footnotes.append(self.current_footnote)
            self.parsing_footnote = False
            self.append_to_previous = False
        elif letter == '‹':
            self.parsing_footnote = True
            self.append_to_previous = False
            self.current_footnote = Footnote(FootnoteType.SYMBOL, None)  # TODO: Get specific type from BS
        elif letter == '›':
            self.parsing_footnote = False
            if self.current_footnote:  # TODO: handle errors
                self.current_footnote.anchor = self.word
                self.word.footnotes.append(self.current_footnote)
                self.current_footnote = None
        else:
            if letter not in ["\uF0A4",""]:  # don't add gibberish
                self.current_footnote.text += letter

    def process_image(self, elem, return_text=False):
        img_src = elem.attrs['src']
        img_location = img_src[img_src.index('image/'):]
        try:
            with open(img_location, 'rb') as img_file:
                img_data = img_file.read()
                b64_img_data = b64encode(img_data)
                elem['src'] = "data:image/{};base64,{}".format('jpg', str(b64_img_data)[2:-1])
                if return_text:
                    return str(elem)
                else:
                    self.process_words(str(elem), None, True)
        except Exception as e:
            print('Exception: {}. Original Source: {}'.format(e, img_src))
            if return_text:
                return ''

    def process_paragraph_elem(self, elem, paragraph, formatting=None):
        #try:
            if isinstance(elem, str):
                self.process_words(elem, formatting)
            elif elem.name == 'span' and 'id' in elem.attrs and 'endnote-'in elem['id']:
                # TODO: figure out where these endnotes are
                pass
            elif elem.name == 'span' and any(x in HtmlTzParser.FOOTNOTES for x in elem['class']):  # other Footnotes
                # footnote_type = [x for x in elem['class'] if x in HtmlTzParser.FOOTNOTES]
                self.process_footnote_material(elem)
            elif elem.name == 'span' and all(x not in HtmlTzParser.FOOTNOTES for x in elem['class']):  # Formatted text
                for child in elem.children:  # exception for gray-text override-9
                    format_class = HtmlTzParser.FORMATTING_CLASSES[[x for x in elem['class'] if x in HtmlTzParser.FORMATTING_CLASSES][0]] if \
                        any(x in HtmlTzParser.FORMATTING_CLASSES for x in elem['class']) else None
                    if len([x for x in elem['class'] if x in HtmlTzParser.APPEND_TO_WORD]) > 0:
                        self.continue_phrase = True
                    self.process_paragraph_elem(child, paragraph, format_class)
            elif elem.name == 'br':
                self.line = self.paragraph.add_new_line()
                self.lines.append(self.line)
                self.append_to_previous = False
            elif elem.name == 'img':
                self.process_image(elem)
            else:
                raise Exception("Unhandled: " + str(elem))
                # TODO: handle unhandled
        #except Exception as e:
        #    logging.error("Error for parsing element " + str(elem) + ". " + str(e))
        # elif #citation
        #     pass
        # new line

    def process_paragraph(self):
        # print([a for a in self.processed_elem_cursor.children])
        self.create_new_paragraph()
        for child in self.processed_elem_cursor.children:
            self.process_paragraph_elem(child, self.paragraph)

    def get_next_word(self):
        return None

    def clean_he(self, hebrew_text):
        if isinstance(hebrew_text, str):
            return hebrew_text
        else:
            cleaned_hebrew_text = ''
            for child in hebrew_text:
                if isinstance(child, str):
                    cleaned_hebrew_text += child
                elif child.name == 'br':
                    cleaned_hebrew_text += "</br>"
                elif child.name == 'span':
                    if 'hebrew-bd' in child.attrs['class']:
                        child.name = 'b'
                        child.attrs = {}
                    else:
                        for child_of_child in child.children:
                            if isinstance(child, str):
                                cleaned_hebrew_text += child_of_child
                            elif child_of_child.name == 'img':
                                cleaned_hebrew_text += self.process_image(child_of_child, True)
                            else:
                                print(str(child))
                    cleaned_hebrew_text += str(child)
                elif child.name == 'img':
                    cleaned_hebrew_text += self.process_image(child, True)
                else:
                    print(str(child))
            return cleaned_hebrew_text

    def parse_hebrew_contents(self):
        hebrew = self.doc_rep.find(id="_idContainer2165")
        types = {}
        daf_index = 0
        tikkun_index = 0
        paragraph_index = 0

        for raw_child in hebrew.children:
            if isinstance(raw_child, str):
                child = raw_child
            else:
                child = self.get_processed_elem(raw_child)
            if isinstance(child, str):
                pass
            if child.name == 'div' and 'daf-hebrew' in child.attrs['class']:  # daf
                if paragraph_index < len(self.daf.paragraphs):
                    print(paragraph_index)
                    print(len(self.daf.paragraphs))
                    print(self.daf.name)
                    print("^paragraph_index, self.daf.paragraphs length, daf -- not enough hebrew")
                self.daf = self.dapim[daf_index]
                daf = child.find_next('p').text
                self.daf.he_name = daf
                daf_index += 1
                paragraph_index = 0
            elif child.name == 'p' and 'chapter-number-title-Hebrew' in child.attrs['class']:
                self.tikkun = self.tikkunim[tikkun_index]
                self.tikkun.he_name = "".join([str(content) for content in child.contents])
                tikkun_index += 1
            elif child.name == 'p' and 'verse-Hebrew' in child.attrs['class']:
                if paragraph_index < len(self.daf.paragraphs):
                    self.paragraph = self.daf.paragraphs[paragraph_index]
                # TODO: Debug misaligned verses
                else:  # somehow misaligned??
                    self.paragraph = Paragraph(self.tikkun, self.daf, next(self.paragraph_number))
                    self.daf.paragraphs.append(self.paragraph)
                    print(self.daf.name)
                    print("too much hebrew")
                self.paragraph.he_words = self.clean_he(child.contents)
                    #''.join([str(content) for content in child.contents])
                paragraph_index += 1

#
# parser2 = DocsTzParser("vol3.docx", 3)
# parser2.read_file()

# print(parser2.doc_rep)

# parser = HtmlTzParser("vol2.html", 1)
# parser.read_file()
# for line in parser2.lines:
#     print("---")
#     for phrase in line.phrases:
#         #if any(['‘' in word.text for word in phrase.words]):
#             print(phrase.formatting)
#             print([word.text for word in phrase.words])

    # for quoted in line.quoted:
    #     print([word.text for word in quoted.words])
# for phrase in parser.phrases:
#     print([word.text for word in phrase.words])
#
# if '‘' in word:
#     if re.match(r'^[?./‘\[\].,:!]+[a-zA-Z0-9]+.*$', str(word)):  # starts with ‘
#         self.line.add_new_quoted()  # create a Quote()
#         if re.match(r'[?.‘\[\].,:!]+$', str(word)):  #
#             self.line.commit_quoted()
#         number += 1
#         # if number != 1:
#         #    print("help!")
#         # self.line.add_new_quoted()
#     elif re.match(r'^[a-zA-Z0-9]+[?.‘\[\].,:!]+$', str(word)):  # or re.match('^[?.‘.:!]+'):
#         self.line.commit_quoted()
#         # self.line.commit_quoted()
#         # print('end quote')
#         # print(word)
#         number -= 1
#         # if number != 0:
#         #    print("help")
#     else:
#         if self.line.inside_quotes:
#             self.line.commit_quoted()
#         else:
#             self.line.add_new_quoted()
#         # pass
#         # print(word)
#     # if it's start
#     # if it's end
#     # if it's both

def run_parse():
    # parsers = [HtmlTzParser("vol2.html", 2, language="bi")]
    parsers = [DocsTzParser("vol3.docx", 3, language="bi")]
    # parsers = [HtmlTzParser("vol2.html", 2), DocsTzParser("vol3.docx", 3), DocsTzParser("vol4.docx", 4), DocsTzParser("vol5.docx", 5)]
    # parsers = [DocsTzParser("vol4.docx", 4)]
    # parsers = [DocsTzParser("vol4.docx", 4), DocsTzParser("vol5.docx", 5)]
    for i, parser in enumerate(parsers):
        if i > 0:
            parser.tikkun = parsers[i-1].tikkun
        parser.read_file()
        with open(f'tz_en_vol{parser.volume}.csv', 'w+') as tz_en:
            fieldnames = ["tikkun", "daf", "paragraph", "he", "line", "phrase", "formatting", "footnotes"]
            csv_writer = csv.DictWriter(tz_en, fieldnames)
            csv_writer.writeheader()
            for paragraph in parser.paragraphs:
                if len(paragraph.phrases) == 0:
                    row = {
                        "tikkun": ' '.join([word.text for word in paragraph.tikkun.words if word.text is not ' ']),
                        "daf": phrase.daf.name,
                        "he": paragraph.he_words
                    }
                else:
                    for phrase in paragraph.phrases:
                        footnotes = []
                        for word in phrase.words:
                            for footnote in word.footnotes:
                                if type(footnote.anchor) is Word:
                                    anchor = footnote.anchor.text
                                else: # Phrase
                                    anchor = ' '.join([word.text for word in footnote.anchor.words])
                                footnotes.append(f"{footnote.footnote_type}/{footnote.formatting}[{anchor}]:{footnote.text}")
                        for footnote in phrase.footnotes:
                            anchor = ' '.join([word.text for word in footnote.anchor.words])
                            footnotes.append(f"{footnote.footnote_type}/{footnote.formatting}[{anchor}]:{footnote.text}")
                        row = {
                            "tikkun": ' '.join([word.text for word in phrase.tikkun.words if word.text is not ' ']),
                            "daf": phrase.daf.name,
                            "paragraph": phrase.paragraph.paragraph_number,
                            "he": phrase.paragraph.he_words,
                            "line": phrase.line.line_number,
                            "phrase": ' '.join([word.text for word in phrase.words]),
                            "formatting": phrase.formatting,
                            "footnotes": ';'.join(footnotes)
                        }
                        csv_writer.writerow(row)

        with open(f'tz_en_vol{parser.volume}_quotes.csv', "w+") as tz_en_quotes:
            fieldnames = ["tikkun", "daf", "paragraph", "line", "he", "quotes", "formatting", "comments"]
            csv_writer = csv.DictWriter(tz_en_quotes, fieldnames)

#run_parse()